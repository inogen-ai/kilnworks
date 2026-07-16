from pathlib import Path

from kilnworks.adapters.media.fake import FakeTranscriber, FakeVisionExtractor
from kilnworks.adapters.sources import parsers
from kilnworks.adapters.sources.parsers import parse_file
from kilnworks.core.ports import MediaExtractor

FIXTURES_DIR = Path(__file__).parent.parent / "fixtures"


def _make_pdf(path, pages):
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    pdf = canvas.Canvas(str(path), pagesize=letter)
    for text in pages:
        pdf.drawString(72, 720, text)
        pdf.showPage()
    pdf.save()


def test_text_metadata_word_count_size_and_type(tmp_path):
    path = tmp_path / "notes.md"
    path.write_text("one two three four five")
    meta = parse_file(path).metadata
    assert meta["word_count"] == 5
    assert meta["size_bytes"] == path.stat().st_size
    assert meta["content_type"] == "text/markdown"


def test_pdf_metadata_page_count(tmp_path):
    path = tmp_path / "doc.pdf"
    _make_pdf(path, ["page one", "page two", "page three"])
    meta = parse_file(path).metadata
    assert meta["page_count"] == 3
    assert meta["content_type"] == "application/pdf"


def test_docx_metadata_word_count(tmp_path):
    import docx

    document = docx.Document()
    document.add_paragraph("alpha beta")
    document.add_paragraph("gamma")
    path = tmp_path / "doc.docx"
    document.save(str(path))
    assert parse_file(path).metadata["word_count"] == 3


def test_csv_metadata_row_count_and_type(tmp_path):
    path = tmp_path / "data.csv"
    path.write_text("name,age\nAlice,30\nBob,25\n")
    meta = parse_file(path).metadata
    assert meta["row_count"] == 3
    assert meta["content_type"] == "text/csv"


def test_xlsx_metadata_sheets_and_rows():
    meta = parse_file(FIXTURES_DIR / "sample.xlsx").metadata
    assert meta["sheet_count"] >= 1
    assert meta["row_count"] >= 1
    assert meta["content_type"].endswith("spreadsheetml.sheet")


def test_image_metadata_dimensions_and_type():
    media = MediaExtractor(vision=FakeVisionExtractor(reply="a square"))
    meta = parse_file(FIXTURES_DIR / "sample.png", media=media).metadata
    assert meta["width"] > 0 and meta["height"] > 0
    assert meta["content_type"] == "image/png"


def test_media_metadata_has_segment_count_and_type():
    media = MediaExtractor(transcription=FakeTranscriber(reply="hello"))
    meta = parse_file(FIXTURES_DIR / "sample.wav", media=media).metadata
    assert isinstance(meta["segment_count"], int)
    assert meta["content_type"] == "audio/wav"


def test_segment_regex_matches_real_transcript_output():
    # Build the text via the real renderer rather than a hand-written literal, so a
    # future change to the transcript format that stops matching _SEGMENT_RE (and
    # silently zeros every media segment_count) fails this test.
    from kilnworks.adapters.media.transcript import render_segments

    text = render_segments([(0.0, "hello"), (12.0, "world"), (90.0, "bye")], "")
    assert len(parsers._SEGMENT_RE.findall(text)) == 3


def test_metadata_capture_is_best_effort(tmp_path, monkeypatch):
    """A wrapped per-format extractor raising must not fail the parse; the
    always-on size/type still land."""

    def boom(*_args):
        raise RuntimeError("nope")

    monkeypatch.setattr(parsers, "_table_metadata", boom)
    path = tmp_path / "data.csv"
    path.write_text("name,age\nAlice,30\n")
    meta = parse_file(path).metadata
    assert "row_count" not in meta  # the failing extractor's key is simply omitted
    assert meta["size_bytes"] > 0  # but the always-on size/type still land
    assert meta["content_type"] == "text/csv"
